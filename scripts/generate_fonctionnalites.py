"""Génère le document Word des fonctionnalités du Portail Socrate."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Couleurs ──────────────────────────────────────────────────────────────────
TEAL      = RGBColor(0x2C, 0x7A, 0x7B)  # couleur primaire du portail
TEAL_DARK = RGBColor(0x1A, 0x52, 0x52)
GREY      = RGBColor(0x44, 0x44, 0x44)
GREY_LIGHT= RGBColor(0x88, 0x88, 0x88)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# ── Marges ────────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def shade_para(para, fill_hex):
    """Fond coloré sur un paragraphe."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)


def add_heading_banner(doc, text, subtitle=None):
    """Bandeau titre teal pleine largeur."""
    # Ligne teal avec titre
    p = doc.add_paragraph()
    shade_para(p, '2C7A7B')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.left_indent  = Cm(0.4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = WHITE
    run.font.name = 'Calibri'
    if subtitle:
        p2 = doc.add_paragraph()
        shade_para(p2, '1A5252')
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(8)
        p2.paragraph_format.left_indent  = Cm(0.4)
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0xB2, 0xD8, 0xD8)
        r2.font.name = 'Calibri'


def h1(doc, text):
    """Titre de section numérotée."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    # Barre latérale teal via bordure gauche
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), '2C7A7B')
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = TEAL_DARK
    run.font.name = 'Calibri'


def h_special(doc, text):
    """Titre de section spéciale (Sécurité, PWA…)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = TEAL
    run.font.name = 'Calibri'
    run.font.italic = True


def bullet(doc, text, bold_prefix=None):
    """Puce standard."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.8)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10)
        r.font.name = 'Calibri'
        r.font.color.rgb = TEAL_DARK
        rest = p.add_run(text)
        rest.font.size = Pt(10)
        rest.font.name = 'Calibri'
        rest.font.color.rgb = GREY
    else:
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.name = 'Calibri'
        r.font.color.rgb = GREY


def note(doc, text):
    """Texte intro/note sous un titre de section."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.4)
    r = p.add_run(text)
    r.font.size = Pt(9.5)
    r.font.color.rgb = GREY_LIGHT
    r.font.italic = True
    r.font.name = 'Calibri'


def separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bot)
    pPr.append(pBdr)


# ─────────────────────────────────────────────────────────────────────────────
# CONTENU
# ─────────────────────────────────────────────────────────────────────────────

# Bandeau principal
add_heading_banner(
    doc,
    'Portail Socrate — Guide des fonctionnalités',
    f'Document destiné au Vénérable Maître  ·  Version Mai 2026',
)

# Introduction
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after  = Pt(6)
r = p.add_run(
    'Le Portail Socrate est la plateforme numérique unifiée de la Loge, '
    'remplaçant les outils dispersés (Telegram, listes mail, Google Drive). '
    'Accessible depuis n\'importe quel navigateur ou installable sur smartphone (application mobile PWA).'
)
r.font.size = Pt(10)
r.font.name = 'Calibri'
r.font.color.rgb = GREY

separator(doc)

# ── Sections ──────────────────────────────────────────────────────────────────

sections_data = [
    ('1. Espace Membres', None, [
        'Annuaire complet des membres avec photo, grade, fonction rituelle',
        'Fiches individuelles : coordonnées, parcours maçonnique, dates d\'initiation / passage / élévation, cotisation',
        'Gestion des civilités F∴ / S∴ automatique',
        'Affichage des responsabilités passées et actuelles',
        'Assiduité individuelle consultable',
        'Groupes & commissions (membres dynamiques)',
    ]),
    ('2. Tenues & Calendrier', None, [
        'Liste des tenues passées et à venir',
        'Inscription / réponse de présence (présent · excusé · absent)',
        'Gestion des visiteurs maçons passants',
        'Gestion des invités pour les agapes',
        'Verrouillage de tenue par le VM',
        'Page Conseil d\'officiers : désignation des remplaçants avant chaque tenue',
        'Page Émargement : feuille de présence en temps réel',
        'Export du récapitulatif de présence',
        'Calendrier général de la Loge',
    ]),
    ('3. Programmes mensuels', None, [
        'Génération automatique du programme de tenue',
        'Envoi par email à tous les membres (avec QR code d\'inscription)',
        'Export PDF joint automatiquement',
        'Possibilité d\'attacher un flyer / affiche',
    ]),
    ('4. Procès-Verbaux (Tracé de tenue)', None, [
        'Rédaction du tracé directement dans le portail',
        'Accès restreint aux officiers (VM, Secrétaire, Surveillants)',
        'Bannière de confidentialité configurable',
        'Stockage sécurisé en base de données',
    ]),
    ('5. Planches', None, [
        'Publication de planches par grade (Apprenti / Compagnon / Maître / Tous)',
        'Commentaires par les membres autorisés',
        'Notification push aux membres concernés lors d\'une nouvelle publication',
        'Téléchargement PDF',
    ]),
    ('6. Forum', None, [
        'Fils de discussion persistants par catégorie',
        'Distinct du Chat (messages éphémères) — conversations de fond',
        'Abonnements aux fils suivis',
    ]),
    ('7. Chat en temps réel', None, [
        'Canaux par grade, par fonction, par commission, ou en direct (message privé)',
        'Messages avec pièces jointes, réponses, suppression',
        'Mentions et notifications',
    ]),
    ('8. Messagerie interne', None, [
        'Messages privés entre membres',
        'Pièces jointes depuis la GED',
        'Notification email (courte, sans divulguer le contenu)',
    ]),
    ('9. Listes de diffusion — Email groupé', None, [
        'Listes statiques (membres choisis) ou dynamiques (par grade, fonction, groupe)',
        '3 listes système créées automatiquement : Tous les actifs, Apprentis+Compagnons, Maîtres',
        'Ajout de contacts externes (passants, contacts institutionnels) avec Prénom, Nom, Loge, Orient',
        'Import en masse par fichier CSV',
        'Corps de l\'email en Markdown avec variables personnalisées (prénom, grade…)',
        'Pièces jointes depuis la GED',
        'Envoi individuel garanti — aucune adresse email visible par les autres destinataires',
        'Lien de désinscription RGPD dans chaque email',
        'Historique des campagnes et statut de livraison par destinataire',
    ]),
    ('10. GED — Bibliothèque documentaire', None, [
        'Espaces → Dossiers → Documents avec droits par grade / groupe',
        'Versioning des documents',
        'Partage externe tokenisé (lien sécurisé à durée limitée)',
        'Espace personnel privé par membre',
        'Pièces jointes dans le Chat et la Messagerie depuis la GED',
    ]),
    ('11. Sondages & Votes', None, [
        'Création de sondages avec choix multiples',
        'Votes anonymes ou nominatifs',
        'Résultats en temps réel',
        'Restriction par grade ou groupe',
    ]),
    ('12. Finance & Cotisations', None, [
        'Budget prévisionnel annuel par catégorie de charges',
        'Calcul automatique du barème T1 → T5 depuis le budget',
        'Appel à tranche avec fenêtre temporelle (15/11 → 31/12) et rappels automatiques J-3',
        'Gestion des capitations nationale et régionale',
        'Suivi des paiements par membre',
        'Quitus annuel',
        'Bilan comptable intermédiaire et annuel',
        'Export CSV des retardataires',
        'Journal des transactions',
    ]),
    ('13. Projets & Commissions', 'Gestion de tâches et de chantiers', [
        'Vues : Liste, Kanban (glisser-déposer), Gantt avec zoom semaines / mois / trimestres',
        'Sous-tâches et dépendances entre tâches',
        'Commentaires, jalons, assignation par membre ou groupe',
        'Templates de projets réutilisables',
        'Dashboard global, rappels push J-3, journal d\'activité',
        'Export iCal et PDF',
    ]),
    ('14. Agapes', 'Maître des Banquets', [
        'Fiche dédiée par tenue : liste nominative des couverts (membres + visiteurs maçons)',
        'Compteur total, régimes alimentaires spéciaux',
        'Export Excel pour le restaurateur',
        'Tableau de bord "Mes agapes" pour le Maître des Banquets',
    ]),
    ('15. Anniversaires maçonniques', None, [
        'Rappels automatiques des anniversaires d\'initiation, passage, élévation',
        'Notification push aux membres concernés la veille',
    ]),
    ('16. Notifications Push', 'Application installable sur smartphone et ordinateur', [
        'Notifications pour : nouvelles tenues, messages, tâches J-3, anniversaires, nouvelles planches',
        'Fonctionne même navigateur fermé',
        'Installable depuis Safari (iPhone) ou Chrome (Android/PC)',
        'Page hors-ligne si réseau indisponible',
    ]),
    ('17. Administration', 'Réservé à l\'administrateur technique', None),
]

for title, subtitle, bullets in sections_data:
    h1(doc, title)
    if subtitle:
        note(doc, subtitle)
    if bullets:
        for b in bullets:
            bullet(doc, b)

# Section 17 spéciale (avec sous-sous-items en gras)
admin_items = [
    ('Vue d\'ensemble : ', 'KPIs loge, alertes (disque, backup), journal d\'activité 30 jours'),
    ('Console utilisateurs : ', 'activation / désactivation, droits admin, reset mot de passe'),
    ('Invitations : ', 'génération d\'un lien d\'accès pour un nouveau membre (valide 7 jours)'),
    ('Journal d\'audit : ', 'qui a fait quoi, quand, depuis quelle adresse IP'),
    ('Données : ', 'sauvegardes (liste, téléchargement, déclenchement manuel), VACUUM DB, export RGPD complet d\'un membre'),
    ('Communication : ', 'journal des emails envoyés, outil de test SMTP'),
    ('Configuration : ', 'libellés d\'affichage personnalisables (grades, fonctions, types de tenue) sans toucher au code'),
    ('Bannière globale : ', 'message d\'information / avertissement affiché sur toutes les pages'),
    ('Confidentialité : ', '3 protections opt-in — restriction des pièces jointes, audit des consultations sensibles, bannière "Confidentiel"'),
]
for prefix, rest in admin_items:
    bullet(doc, rest, bold_prefix=prefix)

separator(doc)

# ── Paramètres de la Loge
h1(doc, '18. Paramètres de la Loge')
for b in [
    'Informations de la loge (nom, orient, logo / sceau)',
    'Tableau de bord des officiers rituels',
    'Configuration SMTP (envoi email)',
    'Carnet de contacts externes (passants, institutionnels)',
    'Gestion des accès utilisateurs',
    'Sauvegarde hebdomadaire automatique',
    'Notifications push (activation par appareil)',
]:
    bullet(doc, b)

separator(doc)

# ── Sécurité
h_special(doc, 'Accès & Sécurité')
for b in [
    'Connexion par identifiant + mot de passe avec réinitialisation par email',
    'Permissions par grade et par fonction (chaque page n\'est accessible qu\'aux habilités)',
    'Journal d\'audit de toutes les actions sensibles (connexions, modifications, exports)',
    'Sauvegarde automatique hebdomadaire (archive ZIP téléchargeable)',
]:
    bullet(doc, b)

# ── PWA
h_special(doc, 'Application Mobile (PWA)')
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after  = Pt(4)
p.paragraph_format.left_indent  = Cm(0.4)
r = p.add_run(
    'Le portail est une Progressive Web App : installable depuis n\'importe quel navigateur, '
    'elle offre une expérience d\'application native sans passer par l\'App Store ni le Play Store.'
)
r.font.size = Pt(10)
r.font.name = 'Calibri'
r.font.color.rgb = GREY

# ── Pied de page
separator(doc)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    f'Document généré depuis le Portail Socrate  ·  Loge Socrate Raison et Progrès  ·  {datetime.date.today().strftime("%B %Y")}'
)
r.font.size = Pt(8)
r.font.color.rgb = GREY_LIGHT
r.font.italic = True
r.font.name = 'Calibri'

# ─────────────────────────────────────────────────────────────────────────────
# Sauvegarde
# ─────────────────────────────────────────────────────────────────────────────
out = 'portail-socrate-fonctionnalites.docx'
doc.save(out)
print(f'OK → {out}')
